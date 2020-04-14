from docx import Document
import os
import csv
import re
from library.FileProcessBasic import FileProcessBasic
import util

def check_output_file(output_path, header):
    if not os.path.exists(output_path):
        with open(output_path, "w", encoding="utf_8_sig", newline="") as f:
            w = csv.DictWriter(f, header)
            w.writeheader()

class Record:
    def __init__(self, docx):
        para_62, para_rest = self.locateParagraph(docx)
        GSI_INTE = self.get_GSI_INTE(para_rest)
        GSI_GPR = self.get_GSI_GPR(para_62)
        GSI_LITH = self.get_GSI_LITH(para_rest)
        GSI_WEA = self.get_GSI_WEA(para_rest)
        GSI_STRU = self.get_GSI_STRU(para_rest)
        GSI_STAB = self.get_GSI_STAB(para_rest)
        GSI_DSCR = self.get_GSI_DSCR(para_rest)
        GSI_PSRL = self.get_GSI_PSRL(para_rest)

        table = docx.tables[2]
        GSI_CHAI = self.get_GSI_CHAI(table)
        GSI_FAUL = self.get_GSI_FAUL(table)
        GSI_WATG = self.get_GSI_WATG(table)

        GSI_WATE = self.get_GSI_WATE()

        self.dict = {
            "掌子面桩号": GSI_CHAI,
            "桩号区间": GSI_INTE,
            "地质雷达描述": GSI_GPR,
            "地下水状态描述": GSI_WATE,
            "地下水对应等级": GSI_WATG,
            "岩性": GSI_LITH,
            "风化程度": GSI_WEA,
            "结构构造": GSI_STRU,
            "断层": GSI_FAUL,
            "稳定性": GSI_STAB,
            "设计围岩级别": GSI_DSCR,
            "预报围岩级别": GSI_PSRL
        }

    def locateParagraph(self, docx):
        flag = 0
        para_62 = ""
        para_rest = ""
        for i, p in enumerate(docx.paragraphs):
            if p.text.startswith("6.2") and flag == 0:
                flag = 1
                continue
            if flag == 1:
                if p.text.startswith("7"):
                    flag = 2
                    para_rest += p.text
                elif p.text.startswith("图"):
                    continue
                else:
                    para_62 += p.text
            if flag == 2:
                para_rest += p.text
        return para_62, para_rest

    def get_GSI_CHAI(self, table):
        # Version by BuDi
        # GSI_CHAI = ''
        # for i in range(len(para)):
        #     if (para[i] == '+'):
        #         j = i
        #         while (para[j] != '（'):
        #             j = j - 1
        #         while (para[j + 1] != '～'):
        #             GSI_CHAI = GSI_CHAI + (para[j + 1])
        #             j = j + 1
        #         break

        for i in range(len(table.rows)):
            tmp = list(table.rows[i].cells)
            cols = sorted(set(tmp), key=tmp.index)
            for j in range(len(cols)):
                if cols[j].text == '掌子面桩号' and j < len(cols) - 1:
                    GSI_CHAI = re.sub(u"\\（.*?）", "", cols[j + 1].text)
                    return GSI_CHAI


    def get_GSI_INTE(self, para):
        GSI_INTE = ''
        for i in range(len(para)):
            if para[i] == '+':
                j = i
                while para[j] != '（':
                    j = j - 1
                while para[j + 1] != '）':
                    GSI_INTE = GSI_INTE + (para[j + 1])
                    j = j + 1
                break
        return GSI_INTE

    def get_GSI_GPR(self, para):
        # 地质雷达描述
        GSI_GPR = "无"
        # for i in range(len(para) - 2):
        #     if para[i:i + 3] == "电磁波":
        #         j = i + 3
        #         k = i + 4
        #         times = 3
        #         while para[j] != '，':
        #             GSI_GPR = para[j] + GSI_GPR
        #             j = j - 1
        #         while para[k] != '，' or times != 0:
        #             GSI_GPR = GSI_GPR + para[k]
        #             k = k + 1
        #             if para[k] == '，':
        #                 times = times - 1
        #         break
        try:
            para = para[para.find("电磁波"):]
            para = para[para.find("，") + 1:]
            GSI_GPR = para
            return GSI_GPR
        except:
            return GSI_GPR

    def get_GSI_LITH(self, para):
        GSI_LITH = ""
        for i in range(len(para) - 2):
            if para[i:i + 3] == "岩性为":
                j = i + 3
                while para[j] != '，':
                    GSI_LITH = GSI_LITH + para[j]
                    j = j + 1
                break
        if GSI_LITH == "":
            GSI_LITH = "无"
        return GSI_LITH

    def get_GSI_WEA(self, para):
        GSI_WEA = ""
        for i in range(len(para) - 1):
            if para[i:i + 2] == "风化":
                GSI_WEA = para[i - 1:i + 2]
                if para[i - 1] == "等":
                    GSI_WEA = para[i - 2:i + 2]
                if para[i - 2] == "～":
                    if para[i - 3] != "等":
                        GSI_WEA = para[i - 3:i + 2]
                    else:
                        GSI_WEA = para[i - 4:i + 2]
                if para[i - 3] == "～":
                    GSI_WEA = para[i - 4:i + 2]
        if GSI_WEA == "":
            GSI_WEA = "无"
        return GSI_WEA

    def get_GSI_STRU(self, para):
        GSI_STRU = ""
        for i in range(len(para) - 1):
            if para[i:i + 2] == "结构":
                j = i + 1
                while para[j] != '，':
                    GSI_STRU = para[j] + GSI_STRU
                    j = j - 1
                break
        if GSI_STRU == "":
            GSI_STRU = "无"
        return GSI_STRU

    def get_GSI_STAB(self, para):
        GSI_STAB = ""
        for i in range(len(para) - 2):
            if para[i:i + 3] == "稳定性":
                j = i + 3
                t = 5
                while para[j] != '，' or t != 0:
                    GSI_STAB = para[j] + GSI_STAB
                    j = j - 1
                    if para[j] == '，':
                        t = t - 1
                break
        if GSI_STAB == "":
            GSI_STAB = "无"
        return GSI_STAB

    def get_GSI_DSCR(self, para):
        GSI_DSCR = ""
        for i in range(len(para) - 6):
            if para[i:i + 7] == "设计围岩等级为":
                GSI_DSCR = para[i + 7]
                break
        if GSI_DSCR == "":
            GSI_DSCR = "无"
        return GSI_DSCR

    def get_GSI_PSRL(self, para):
        GSI_PSRL = ""
        for i in range(len(para) - 4):
            if para[i:i + 5] == "预判围岩为":
                GSI_PSRL = para[i + 5]
                break
        if GSI_PSRL == "":
            GSI_PSRL = "无"
        return GSI_PSRL

    def get_GSI_WATE(self):
        return "无"

    def get_GSI_WATG(self, table):
        for i in range(len(table.rows)):
            text = table.cell(i, 0).text
            if text == '地下水状态':
                tmp = list(table.rows[i].cells)
                cols = sorted(set(tmp), key=tmp.index)
                for col in cols:
                    if col.text.find('√') > 0:
                        col.text = col.text.replace('√', '')
                        GSI_WATG = col.text
                        return GSI_WATG

    def get_GSI_FAUL(self, table):
        # version by BuDi
        # GSI_FAUL = '无'
        # for i in range(len(para) - 1):
        #     if (para[i:i + 2] == "断层"):
        #         GSI_FAUL = "断层带"

        GSI_FAUL = ""
        for i in range(len(table.rows)):
            text = table.cell(i, 0).text
            if text == '岩体出露状态':
                tmp = list(table.rows[i].cells)
                cols = sorted(set(tmp), key=tmp.index)
                for col in cols:
                    if col.text.find('√') > 0:
                        col.text = col.text.replace('√', '')
                        GSI_FAUL = col.text
                        return GSI_FAUL
        if GSI_FAUL == "":
            GSI_FAUL = "无"
        return GSI_FAUL

class Processor(FileProcessBasic):
    def save(self, output, record):
        output_path = os.path.join(output, "GPR_S1S2.csv")
        header = record.dict.keys()
        check_output_file(output_path, header)

        with open(output_path, "a+", encoding="utf_8_sig", newline="") as f:
            w = csv.DictWriter(f, record.dict.keys())
            w.writerow(record.dict)

    def run(self, input_path, output_path):
        transformed = False
        if input_path.endswith("doc"):
            transformed = True
            util.doc2docx(input_path)
            input_path = input_path + "x"
        elif not input_path.endswith("docx"):
            return

        docx = Document(input_path)
        record = Record(docx)
        self.save(output_path, record)

        print("提取完成" + input_path)
        if transformed and os.path.exists(input_path):
            os.remove(input_path)


if __name__ == "__main__":
    test = Processor()
    inputpath = "D:/Death in TJU/Junior_2nd/iS3 Lab2/tasks/task3/GPRS1S2.docx"
    outputpath = "D:/Death in TJU/Junior_2nd/iS3 Lab2/tasks/task3"
    test.run(inputpath, outputpath)